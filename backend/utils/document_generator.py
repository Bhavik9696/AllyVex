import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ── PDF imports ──
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.platypus import KeepTogether

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_colored_paragraph(doc, label, value, label_color=(0, 0, 0)):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(*label_color)
    p.add_run(str(value))
    return p

def add_section_divider(doc):
    doc.add_paragraph("─" * 60)

def get_verdict_color(verdict):
    if verdict == "PURSUE":
        return (0, 128, 0)      # Green
    elif verdict == "HOLD":
        return (255, 140, 0)    # Orange
    else:
        return (200, 0, 0)      # Red

def generate_dossier(
    company_name: str,
    domain: str,
    client_info: str,
    bull_output: dict,
    bear_output: dict,
    detective_output: dict,
    orchestrator_output: dict,
    track: str = "customer"   # "customer" or "partner"
) -> str:
    """
    Generate a full DOCX intelligence dossier for the given track.
    track="customer" focuses on the customer sales angle.
    track="partner"  focuses on the partnership angle.
    Returns the saved filename inside the outputs/ directory.
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    is_partner = track.lower() == "partner"

    customer_track = orchestrator_output.get("customerTrack", {}) or {}
    partner_track  = orchestrator_output.get("partnerTrack",  {}) or {}

    # Pick the active track data
    active_track = partner_track if is_partner else customer_track
    track_label  = "PARTNER" if is_partner else "CUSTOMER"

    verdict        = active_track.get("verdict", orchestrator_output.get("verdict", "UNKNOWN"))
    verdict_color  = get_verdict_color(verdict)
    confidence     = active_track.get("confidence", orchestrator_output.get("confidence", "N/A"))
    regret         = active_track.get("regretScore", orchestrator_output.get("regretScore", {})) or {}
    recommended_approach = orchestrator_output.get("recommendedApproach", "")
    generated_date = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # ── Cover Header ──
    title = doc.add_heading("ALLYVEX", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph(f"Autonomous Strategic Sales Intelligence Report — {track_label} TRACK")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    # ── Target Company ──
    target_heading = doc.add_paragraph()
    target_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target_heading.add_run(f"Target Company: {company_name}")
    run.bold = True
    run.font.size = Pt(16)

    domain_para = doc.add_paragraph()
    domain_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    domain_para.add_run(domain).font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")

    # ── Verdict Badge ──
    verdict_para = doc.add_paragraph()
    verdict_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    verdict_run = verdict_para.add_run(f"{track_label} VERDICT: {verdict}")
    verdict_run.bold = True
    verdict_run.font.size = Pt(20)
    verdict_run.font.color.rgb = RGBColor(*verdict_color)

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.add_run(
        f"Confidence: {confidence}/100   |   "
        f"Regret Score: {regret.get('score', 'N/A')}/100"
    ).font.size = Pt(11)

    regret_para = doc.add_paragraph()
    regret_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    regret_para.add_run(
        f"{regret.get('reason', '')}"
    ).font.color.rgb = RGBColor(80, 80, 80)

    # Show the other track's verdict for context
    other_track  = customer_track if is_partner else partner_track
    other_label  = "CUSTOMER" if is_partner else "PARTNER"
    other_verdict = other_track.get("verdict", "N/A")
    if other_verdict and other_verdict != "N/A":
        other_para = doc.add_paragraph()
        other_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        other_para.add_run(
            f"(For reference — {other_label} Track Verdict: {other_verdict})"
        ).font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph("")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(
        f"Generated by ALLYVEX on {generated_date}"
    ).font.color.rgb = RGBColor(150, 150, 150)

    add_section_divider(doc)
    doc.add_page_break()

    # ════════════════════════════════════════
    # SECTION 1: EXECUTIVE SUMMARY
    # ════════════════════════════════════════
    add_heading(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        orchestrator_output.get("executiveSummary", "No summary available.")
    )

    doc.add_paragraph("")
    add_colored_paragraph(doc, f"{track_label} Track Verdict", verdict, label_color=verdict_color)
    add_colored_paragraph(doc, "Confidence Level", f"{confidence}/100")
    add_colored_paragraph(
        doc, "Regret Score",
        f"{regret.get('score', 'N/A')}/100 — {regret.get('reason', '')}"
    )
    if recommended_approach:
        add_colored_paragraph(doc, "Recommended Approach", recommended_approach)

    # ════════════════════════════════════════
    # SECTION 2: CLIENT ADVANTAGES & DISADVANTAGES
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "2. Client Position Analysis", level=1)

    add_heading(doc, "Advantages Our Client Holds", level=2)
    advantages = orchestrator_output.get("clientAdvantages", [])
    for adv in advantages:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(adv)

    doc.add_paragraph("")
    add_heading(doc, "Disadvantages & Obstacles", level=2)
    disadvantages = orchestrator_output.get("clientDisadvantages", [])
    for dis in disadvantages:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(dis)

    # ════════════════════════════════════════
    # SECTION 3: DECIDING FACTORS
    # ════════════════════════════════════════
    doc.add_paragraph("")
    add_heading(doc, "3. Key Deciding Factors", level=1)

    active_factors = active_track.get("decidingFactors", {}) or {}
    # Customer track also checks top-level decidingFactors as fallback
    if not active_factors and not is_partner:
        active_factors = orchestrator_output.get("decidingFactors", {}) or {}

    add_heading(doc, f"{track_label} Track — Key Swing Factor", level=2)
    doc.add_paragraph(active_factors.get("keySwingFactor", "No swing factor identified."))

    add_heading(doc, f"{track_label} Track — Strongest Signal FOR", level=2)
    doc.add_paragraph(active_factors.get("strongestBullSignal", "N/A"))

    add_heading(doc, f"{track_label} Track — Strongest Signal AGAINST", level=2)
    doc.add_paragraph(active_factors.get("strongestBearSignal", "N/A"))

    if is_partner:
        add_heading(doc, "Distribution Value Assessment", level=2)
        doc.add_paragraph(active_factors.get("distributionValueVerdict", "N/A"))
    else:
        add_heading(doc, "Scale Verdict", level=2)
        doc.add_paragraph(active_factors.get("scaleVerdict", "N/A"))

    add_heading(doc, "Detective Impact", level=2)
    doc.add_paragraph(active_factors.get("detectiveImpact", "N/A"))

    if recommended_approach:
        add_heading(doc, "Recommended Approach", level=2)
        approach_para = doc.add_paragraph(recommended_approach)
        approach_para.runs[0].bold = True
        doc.add_paragraph(orchestrator_output.get("recommendedApproachReason", ""))

    # ════════════════════════════════════════
    # SECTION 4: BULL AGENT FINDINGS
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, f"4. Bull Agent Findings — The Case FOR {track_label} Pursuit", level=1)

    bull_score = bull_output.get("overallBullScore", "N/A")
    bull_score_display = (
        f"Customer: {bull_score.get('customerScore')}/100, Partner: {bull_score.get('partnerScore')}/100"
        if isinstance(bull_score, dict)
        else f"{bull_score}/100"
    )
    key_arg = bull_output.get("keyArgument", {})
    key_arg_field = "asPartner" if is_partner else "asCustomer"
    key_arg_text = key_arg.get(key_arg_field, str(key_arg)) if isinstance(key_arg, dict) else str(key_arg)
    doc.add_paragraph(f"Bull Score: {bull_score_display} — {key_arg_text}")

    if is_partner:
        # Partner track: show partner signals prominently
        add_heading(doc, "Partner Signals", level=2)
        for signal in bull_output.get("partnerSignals", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{signal.get('strength', '')}] {signal.get('signal', '')}")
            run.bold = True
            p.add_run(f"\n    Partner value: {signal.get('partnerConnection', '')}")
            p.add_run(f"\n    Partnership type: {signal.get('partnerType', 'N/A')}")
            p.add_run(f"\n    Scale reasoning: {signal.get('scaleDependentReasoning', '')}")
            p.add_run(f"\n    Source: {signal.get('source', 'N/A')}")

        add_heading(doc, "Customer Buying Signals (context)", level=2)
        for signal in bull_output.get("customerSignals", bull_output.get("clientRelevantSignals", []))[:3]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{signal.get('strength', '')}] {signal.get('signal', '')}")
            run.bold = True
            p.add_run(f"\n    Why it matters: {signal.get('clientConnection', '')}")
            p.add_run(f"\n    Source: {signal.get('source', 'N/A')}")
    else:
        # Customer track: show customer signals prominently
        add_heading(doc, "Customer Buying Signals", level=2)
        for signal in bull_output.get("customerSignals", bull_output.get("clientRelevantSignals", [])):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{signal.get('strength', '')}] {signal.get('signal', '')}")
            run.bold = True
            p.add_run(f"\n    Why it matters: {signal.get('clientConnection', '')}")
            p.add_run(f"\n    Source: {signal.get('source', 'N/A')}")

        add_heading(doc, "Partner Signals (context)", level=2)
        for signal in bull_output.get("partnerSignals", [])[:3]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{signal.get('strength', '')}] {signal.get('signal', '')}")
            run.bold = True
            p.add_run(f"\n    Partner value: {signal.get('partnerConnection', '')}")
            p.add_run(f"\n    Source: {signal.get('source', 'N/A')}")

    # Technical debt signals
    tech_signals = bull_output.get("technicalDebtSignals", [])
    if tech_signals:
        add_heading(doc, "Technical Debt Opportunities", level=2)
        for ts in tech_signals:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ts.get("observation", "")).bold = True
            p.add_run(f"\n    How our client helps: {ts.get('howClientHelps', '')}")

    # Fiscal pressure signals
    fiscal_signals = bull_output.get("fiscalPressureSignals", [])
    if fiscal_signals:
        add_heading(doc, "Fiscal Pressure Opportunities", level=2)
        for fs in fiscal_signals:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(fs.get("observation", "")).bold = True
            p.add_run(f"\n    How our client helps: {fs.get('howClientHelps', '')}")

    # Pivot signals
    pivot_signals = bull_output.get("recentPivotSignals", [])
    if pivot_signals:
        add_heading(doc, "Recent Pivot Opportunities", level=2)
        for ps in pivot_signals:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ps.get("observation", "")).bold = True
            p.add_run(
                f"\n    New requirement created: {ps.get('newRequirement', '')}"
            )
            p.add_run(f"\n    Client fit: {ps.get('clientFit', '')}")

    # ════════════════════════════════════════
    # SECTION 5: BEAR AGENT FINDINGS
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "5. Bear Agent Findings — The Case AGAINST Pursuing", level=1)

    # ════════════════════════════════════════
    # SECTION 5: BEAR AGENT FINDINGS
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, f"5. Bear Agent Findings — The Case AGAINST {track_label} Pursuit", level=1)

    bear_score = bear_output.get("overallBearScore", "N/A")
    bear_score_display = (
        f"Customer: {bear_score.get('customerScore')}/100, Partner: {bear_score.get('partnerScore')}/100"
        if isinstance(bear_score, dict)
        else f"{bear_score}/100"
    )
    bear_key_arg = bear_output.get("keyArgument", {})
    bear_key_arg_field = "asPartner" if is_partner else "asCustomer"
    bear_key_arg_text = bear_key_arg.get(bear_key_arg_field, str(bear_key_arg)) if isinstance(bear_key_arg, dict) else str(bear_key_arg)
    doc.add_paragraph(f"Bear Score: {bear_score_display} — {bear_key_arg_text}")

    # Deal killers — show the relevant track's deal killer prominently
    deal_killer = bear_output.get("dealKiller")
    if deal_killer and isinstance(deal_killer, dict):
        dk_key = "partnerDealKiller" if is_partner else "customerDealKiller"
        dk_text = deal_killer.get(dk_key)
        if dk_text and str(dk_text) not in ("null", "None", ""):
            p = doc.add_paragraph()
            run = p.add_run(f"[{track_label}] DEAL KILLER: {dk_text}")
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)
    elif deal_killer and str(deal_killer) not in ("null", "None", ""):
        p = doc.add_paragraph()
        run = p.add_run(f"DEAL KILLER: {deal_killer}")
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)

    if is_partner:
        # Partner track: show partner red flags prominently
        add_heading(doc, "Partner Red Flags", level=2)
        for flag in bear_output.get("partnerRedFlags", []):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag.get('severity', '')}] {flag.get('flag', '')}")
            run.bold = True
            p.add_run(f"\n    Partnership impact: {flag.get('partnershipImpact', '')}")
            p.add_run(f"\n    Deal-breaking potential: {flag.get('dealBreakingPotential', 'N/A')}")
            p.add_run(f"\n    Source: {flag.get('source', 'N/A')}")

        add_heading(doc, "Customer Red Flags (context)", level=2)
        for flag in bear_output.get("customerRedFlags", bear_output.get("clientRelevantRedFlags", []))[:3]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag.get('severity', '')}] {flag.get('flag', '')}")
            run.bold = True
            p.add_run(f"\n    Impact: {flag.get('clientImpact', '')}")
            p.add_run(f"\n    Source: {flag.get('source', 'N/A')}")
    else:
        # Customer track: show customer red flags prominently
        add_heading(doc, "Customer Red Flags", level=2)
        for flag in bear_output.get("customerRedFlags", bear_output.get("clientRelevantRedFlags", [])):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag.get('severity', '')}] {flag.get('flag', '')}")
            run.bold = True
            p.add_run(f"\n    Impact on client: {flag.get('clientImpact', '')}")
            p.add_run(f"\n    Source: {flag.get('source', 'N/A')}")

        add_heading(doc, "Partner Red Flags (context)", level=2)
        for flag in bear_output.get("partnerRedFlags", [])[:3]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{flag.get('severity', '')}] {flag.get('flag', '')}")
            run.bold = True
            p.add_run(f"\n    Partnership impact: {flag.get('partnershipImpact', '')}")
            p.add_run(f"\n    Source: {flag.get('source', 'N/A')}")

    # ════════════════════════════════════════
    # SECTION 6: DETECTIVE AUDIT
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "6. Detective Agent Audit — Evidence Quality Review", level=1)

    # Support both new and old schema keys
    customer_audit = detective_output.get("customerTrackAudit", detective_output.get("bullAudit", {})) or {}
    partner_audit = detective_output.get("partnerTrackAudit", detective_output.get("bearAudit", {})) or {}
    bull_audit = detective_output.get("bullAudit", {}) or {}
    bear_audit = detective_output.get("bearAudit", {}) or {}
    scale_ver = detective_output.get("scaleVerification", {}) or {}

    add_heading(doc, "Evidence Quality Scores", level=2)
    add_colored_paragraph(
        doc, "Customer Track Evidence Score",
        f"{customer_audit.get('evidenceScore', bull_audit.get('evidenceScore', 'N/A'))}/100"
    )
    add_colored_paragraph(
        doc, "Partner Track Evidence Score",
        f"{partner_audit.get('evidenceScore', bear_audit.get('evidenceScore', 'N/A'))}/100"
    )
    add_colored_paragraph(
        doc, "Overall Debate Confidence",
        f"{detective_output.get('overallConfidenceInDebate', 'N/A')}/100"
    )

    if scale_ver:
        add_heading(doc, "Scale Verification", level=2)
        add_colored_paragraph(doc, "Confirmed Scale", scale_ver.get("confirmedScale", "N/A"))
        add_colored_paragraph(doc, "Customer Impact", scale_ver.get("scaleImpactsCustomerTrack", "N/A"))
        add_colored_paragraph(doc, "Partner Impact", scale_ver.get("scaleImpactsPartnerTrack", "N/A"))

    split = detective_output.get("splitVerdictAssessment", {}) or {}
    if split:
        add_heading(doc, "Split Verdict Assessment", level=2)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        rows_data = [
            ("Customer Track Strength", split.get("customerTrackStrength", "N/A")),
            ("Partner Track Strength", split.get("partnerTrackStrength", "N/A")),
            ("Recommended Approach", split.get("recommendedApproach", "N/A")),
        ]
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i].cells
            row[0].text = label
            row[1].text = value
        doc.add_paragraph("")
        doc.add_paragraph(split.get("splitReasoning", ""))
    else:
        # Fallback to old clientFitAssessment structure
        client_fit = detective_output.get("clientFitAssessment", {}) or {}
        add_heading(doc, "Client Fit Assessment", level=2)
        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Dimension"
        headers[1].text = "Rating"
        headers[2].text = "Reason"
        for cell in headers:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        fit_rows = [
            ("Technical Fit", customer_audit.get("technicalFit", client_fit.get("technicalFit", "N/A")),
             customer_audit.get("technicalFitReason", client_fit.get("technicalFitReason", ""))),
            ("Budget Fit", customer_audit.get("budgetFit", client_fit.get("budgetFit", "N/A")),
             customer_audit.get("budgetFitReason", client_fit.get("budgetFitReason", ""))),
            ("Timing Fit", customer_audit.get("timingFit", client_fit.get("timingFit", "N/A")),
             customer_audit.get("timingFitReason", client_fit.get("timingFitReason", "")))
        ]
        for i, (dim, rating, reason) in enumerate(fit_rows, 1):
            row = table.rows[i].cells
            row[0].text = dim
            row[1].text = rating
            row[2].text = reason

    add_heading(doc, "Critical Overlooked Fact", level=2)
    doc.add_paragraph(
        detective_output.get("criticalOverlookedFact", "None identified.")
    )

    missing = detective_output.get("missingContext", [])
    if missing:
        add_heading(doc, "Missing Context Found", level=2)
        for mc in missing:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(mc.get("finding", "")).bold = True
            p.add_run(
                f"\n    Impact: {mc.get('impact', '')} — "
                f"{mc.get('clientRelevance', mc.get('explanation', ''))}"
            )

    # ════════════════════════════════════════
    # SECTION 7: OUTREACH STRATEGY
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, f"7. {track_label} Track Outreach Strategy", level=1)

    active_dm = active_track.get("targetDecisionMaker", {}) or {}
    if not active_dm and not is_partner:
        active_dm = orchestrator_output.get("targetDecisionMaker", {}) or {}

    add_heading(doc, "Target Decision Maker", level=2)
    add_colored_paragraph(doc, "Title", active_dm.get("title", "N/A"))
    add_colored_paragraph(doc, "Why This Person", active_dm.get("why", "N/A"))
    add_colored_paragraph(doc, "LinkedIn Search", active_dm.get("linkedinSearchTip", "N/A"))

    active_email = active_track.get("outreachEmail", {}) or {}
    if not active_email and not is_partner:
        active_email = orchestrator_output.get("outreachEmail", {}) or {}

    add_heading(doc, "Outreach Email Draft", level=2)
    add_colored_paragraph(doc, "Subject", active_email.get("subject", "N/A"))
    doc.add_paragraph("")
    doc.add_paragraph(active_email.get("body", "No email generated."))

    add_heading(doc, "Proposed Next Steps", level=2)
    for step in orchestrator_output.get("proposedNextSteps", []):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    # Hold/Avoid conditions for the active track
    if verdict == "HOLD":
        hold_text = active_track.get("ifHold", orchestrator_output.get("ifHold"))
        if hold_text:
            add_heading(doc, "Re-evaluation Trigger", level=2)
            doc.add_paragraph(hold_text)

    if verdict == "AVOID":
        avoid_text = active_track.get("ifAvoid", orchestrator_output.get("ifAvoid"))
        if avoid_text:
            add_heading(doc, "Conditions to Revisit", level=2)
            doc.add_paragraph(avoid_text)

    # ════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════
    add_section_divider(doc)
    footer_para = doc.add_paragraph(
        f"This report was generated autonomously by ALLYVEX on {generated_date}. "
        f"All intelligence was gathered from public sources via live web search. "
        f"This document is confidential and intended for internal sales use only."
    )
    footer_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    footer_para.runs[0].font.size = Pt(9)

    # ── Save ──
    _outputs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs")
    os.makedirs(_outputs_dir, exist_ok=True)
    filename = (
        f"ALLYVEX_{company_name.replace(' ', '_')}_{track_label}_Dossier_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    filepath = os.path.join(_outputs_dir, filename)
    doc.save(filepath)
    print(f"  [DOCUMENT] {track_label} Dossier saved: {filepath}")
    return filename


# ════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY PDF
# ════════════════════════════════════════════════════════════════════

def _get_verdict_rgb(verdict: str):
    """Return a reportlab Color for the given verdict string."""
    if verdict == "PURSUE":
        return colors.Color(0, 0.502, 0)       # Green
    elif verdict == "HOLD":
        return colors.Color(1, 0.549, 0)        # Orange
    else:
        return colors.Color(0.784, 0, 0)        # Red


def generate_executive_summary_pdf(
    company_name: str,
    domain: str,
    orchestrator_output: dict,
    track: str = "customer"   # "customer" or "partner"
) -> str:
    """
    Generate a polished Executive Summary PDF for the given track.
    Returns the saved filename inside the outputs/ directory.
    """
    is_partner = track.lower() == "partner"
    track_label = "PARTNER" if is_partner else "CUSTOMER"

    _outputs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs")
    os.makedirs(_outputs_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"ALLYVEX_{company_name.replace(' ', '_')}_{track_label}_ExecutiveSummary_{timestamp}.pdf"
    filepath = os.path.join(_outputs_dir, filename)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
        title=f"ALLYVEX {track_label} Executive Summary — {company_name}",
        author="ALLYVEX",
    )

    # ── Colour palette ──────────────────────────────────────────────
    DARK        = colors.Color(0.118, 0.118, 0.118)
    MID_GREY    = colors.Color(0.392, 0.392, 0.392)
    LIGHT_GREY  = colors.Color(0.588, 0.588, 0.588)
    RULE_GREY   = colors.Color(0.85,  0.85,  0.85)

    customer_track = orchestrator_output.get("customerTrack", {}) or {}
    partner_track  = orchestrator_output.get("partnerTrack",  {}) or {}
    active_track   = partner_track if is_partner else customer_track
    other_track    = customer_track if is_partner else partner_track
    other_label    = "CUSTOMER" if is_partner else "PARTNER"

    verdict       = active_track.get("verdict", orchestrator_output.get("verdict", "UNKNOWN"))
    verdict_color = _get_verdict_rgb(verdict)
    confidence    = active_track.get("confidence", orchestrator_output.get("confidence", "N/A"))
    regret        = active_track.get("regretScore", orchestrator_output.get("regretScore", {})) or {}
    recommended   = orchestrator_output.get("recommendedApproach", "")
    summary_text  = orchestrator_output.get("executiveSummary", "No summary available.")
    generated_date = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # ── Styles ──────────────────────────────────────────────────────
    def _style(name, **kw):
        return ParagraphStyle(name, **kw)

    s_brand = _style(f"Brand_{track_label}",
        fontSize=26, leading=30, textColor=DARK,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=2)

    s_sub = _style(f"Sub_{track_label}",
        fontSize=11, leading=14, textColor=MID_GREY,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=4)

    s_target = _style(f"Target_{track_label}",
        fontSize=15, leading=18, textColor=DARK,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=2)

    s_domain = _style(f"Domain_{track_label}",
        fontSize=10, leading=13, textColor=LIGHT_GREY,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=12)

    s_verdict = _style(f"Verdict_{track_label}",
        fontSize=20, leading=24, textColor=verdict_color,
        fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)

    s_conf = _style(f"Conf_{track_label}",
        fontSize=10, leading=13, textColor=MID_GREY,
        fontName="Helvetica", alignment=TA_CENTER, spaceAfter=4)

    s_other_verdict = _style(f"OtherVerdict_{track_label}",
        fontSize=10, leading=13, textColor=MID_GREY,
        fontName="Helvetica-Oblique", alignment=TA_CENTER, spaceAfter=4)

    s_regret_reason = _style(f"RegretReason_{track_label}",
        fontSize=9, leading=12, textColor=LIGHT_GREY,
        fontName="Helvetica-Oblique", alignment=TA_CENTER, spaceAfter=14)

    s_section_head = _style(f"SectionHead_{track_label}",
        fontSize=12, leading=15, textColor=DARK,
        fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=5)

    s_body = _style(f"Body_{track_label}",
        fontSize=10, leading=14, textColor=DARK,
        fontName="Helvetica", spaceAfter=6)

    s_footer = _style(f"Footer_{track_label}",
        fontSize=8, leading=11, textColor=LIGHT_GREY,
        fontName="Helvetica", alignment=TA_CENTER, spaceBefore=12)

    # ── Story ────────────────────────────────────────────────────────
    story = []

    # Brand header
    story.append(Paragraph("ALLYVEX", s_brand))
    story.append(Paragraph(f"Autonomous Strategic Sales Intelligence — {track_label} TRACK", s_sub))
    story.append(HRFlowable(width="100%", thickness=1, color=RULE_GREY, spaceAfter=10))

    # Company info
    story.append(Paragraph(f"Target Company: {company_name}", s_target))
    story.append(Paragraph(domain, s_domain))

    # Active track verdict
    story.append(Paragraph(f"{track_label} VERDICT: {verdict}", s_verdict))
    story.append(Paragraph(
        f"Confidence: {confidence}/100 &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"Regret Score: {regret.get('score', 'N/A')}/100",
        s_conf))
    if regret.get("reason"):
        story.append(Paragraph(regret["reason"], s_regret_reason))

    # Other track verdict for context
    other_verdict = other_track.get("verdict", "N/A")
    if other_verdict and other_verdict != "N/A":
        story.append(Paragraph(
            f"(For reference — {other_label} Track Verdict: {other_verdict})",
            s_other_verdict))

    if recommended:
        story.append(Paragraph(f"Recommended Approach: {recommended}", s_conf))

    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE_GREY, spaceAfter=8))

    # Executive summary body
    story.append(Paragraph("Executive Summary", s_section_head))
    story.append(Paragraph(summary_text, s_body))

    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE_GREY, spaceAfter=8))

    # Key metrics table
    story.append(Paragraph(f"{track_label} Track — Key Metrics", s_section_head))

    other_confidence = other_track.get("confidence", "N/A")
    other_regret     = other_track.get("regretScore", {}) or {}

    metrics_data = [
        ["Metric", "Value"],
        [f"{track_label} Verdict",    verdict],
        [f"{track_label} Confidence", f"{confidence}/100"],
        [f"{track_label} Regret",     f"{regret.get('score', 'N/A')}/100"],
        [f"{other_label} Verdict",    other_verdict],
        [f"{other_label} Confidence", f"{other_confidence}/100"],
        [f"{other_label} Regret",     f"{other_regret.get('score', 'N/A')}/100"],
        ["Recommended Approach",      recommended or "N/A"],
    ]

    tbl = Table(metrics_data, colWidths=[2.5 * inch, 4.0 * inch])
    other_verdict_color = _get_verdict_rgb(other_verdict)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), colors.Color(0.93, 0.93, 0.93)),
        ("TEXTCOLOR",    (0, 0), (-1, 0), DARK),
        ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, 0), 9),
        ("FONTNAME",     (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE",     (0, 1), (-1, -1), 9),
        ("TEXTCOLOR",    (0, 1), (0, -1), MID_GREY),
        ("TEXTCOLOR",    (1, 1), (1, -1), DARK),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.Color(0.97, 0.97, 0.97)]),
        ("GRID",         (0, 0), (-1, -1), 0.5, RULE_GREY),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        # Active track verdict row coloured
        ("TEXTCOLOR",    (1, 1), (1, 1), verdict_color),
        ("FONTNAME",     (1, 1), (1, 1), "Helvetica-Bold"),
        # Other track verdict row coloured
        ("TEXTCOLOR",    (1, 4), (1, 4), other_verdict_color),
        ("FONTNAME",     (1, 4), (1, 4), "Helvetica-Bold"),
    ]))
    story.append(tbl)

    # Footer
    story.append(HRFlowable(width="100%", thickness=1, color=RULE_GREY, spaceBefore=14, spaceAfter=4))
    story.append(Paragraph(
        f"Generated autonomously by ALLYVEX on {generated_date}. "
        "All intelligence was gathered from public sources via live web search. "
        "Confidential — for internal sales use only.",
        s_footer,
    ))

    doc.build(story)
    print(f"  [PDF] {track_label} Executive Summary saved: {filepath}")
    return filename